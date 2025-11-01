#!/usr/bin/env python3
"""
Export Server - Handles PDF/Word export requests
Runs alongside MkDocs dev server
"""

from flask import Flask, send_file, jsonify, send_from_directory
from flask_cors import CORS
import os
import sys
from pathlib import Path
import threading
import time
import subprocess
import markdown
from bs4 import BeautifulSoup

app = Flask(__name__)
CORS(app)  # Enable CORS for MkDocs dev server

EXPORTS_DIR = Path('exports')
EXPORTS_DIR.mkdir(exist_ok=True)

SITE_DIR = Path('site')
DOCS_DIR = Path('docs')

# Cache generated files to avoid regeneration
file_cache = {}


def get_version_html_files(doc_type, version):
    """Get all HTML files for a specific version from site directory"""
    version_path = SITE_DIR / doc_type / 'versions' / version
    if not version_path.exists():
        print(f"Error: Version path {version_path} does not exist")
        return []
    
    html_files = []
    for html_file in sorted(version_path.rglob('index.html')):
        html_files.append(html_file)
    
    print(f"Found {len(html_files)} HTML files for {doc_type} {version}")
    return html_files


def combine_html_for_pdf(doc_type, version):
    """Combine all HTML files into one for PDF generation"""
    html_files = get_version_html_files(doc_type, version)
    if not html_files:
        return None
    
    combined_content = ""
    
    for html_file in html_files:
        print(f"Processing: {html_file}")
        with open(html_file, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
            
            # Extract main content
            article = soup.find('article', class_='md-content__inner')
            if article:
                combined_content += str(article) + "\n\n"
    
    # Create complete HTML with RTL support
    html_template = f"""
<!DOCTYPE html>
<html dir="rtl" lang="fa">
<head>
    <meta charset="UTF-8">
    <title>{doc_type} - {version}</title>
    <style>
        @font-face {{
            font-family: 'Vazirmatn';
            font-weight: 400;
            src: local('Vazirmatn'), local('Vazirmatn-Regular');
        }}
        
        body {{
            font-family: 'Vazirmatn', 'Tahoma', sans-serif;
            direction: rtl;
            text-align: right;
            line-height: 1.8;
            color: #333;
            padding: 40px;
            max-width: 900px;
            margin: 0 auto;
        }}
        
        h1 {{
            color: #1a237e;
            border-bottom: 3px solid #3f51b5;
            padding-bottom: 10px;
            margin-top: 40px;
            page-break-before: always;
        }}
        
        h1:first-child {{
            page-break-before: avoid;
        }}
        
        h2 {{
            color: #283593;
            margin-top: 30px;
            border-right: 4px solid #3f51b5;
            padding-right: 15px;
        }}
        
        h3 {{
            color: #303f9f;
            margin-top: 20px;
        }}
        
        p {{
            margin: 15px 0;
            text-align: justify;
        }}
        
        ul, ol {{
            margin: 15px 0;
            padding-right: 30px;
        }}
        
        li {{
            margin: 8px 0;
        }}
        
        code {{
            background: #f5f5f5;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            direction: ltr;
            display: inline-block;
        }}
        
        pre {{
            background: #f5f5f5;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            direction: ltr;
            text-align: left;
        }}
        
        blockquote {{
            border-right: 4px solid #9fa8da;
            padding: 10px 20px;
            margin: 20px 0;
            background: #f5f7ff;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: right;
        }}
        
        th {{
            background: #3f51b5;
            color: white;
            font-weight: 600;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 50px;
            padding: 30px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 10px;
        }}
        
        .header h1 {{
            color: white;
            border: none;
            margin: 0;
            page-break-before: avoid;
        }}
        
        @page {{
            size: A4;
            margin: 2cm;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>مستندات {doc_type.replace('-', ' ')}</h1>
        <p>نسخه {version}</p>
    </div>
    
    {combined_content}
    
    <div style="margin-top: 50px; padding-top: 20px; border-top: 2px solid #e0e0e0; text-align: center; font-size: 12px; color: #666;">
        <p>این مستند به صورت خودکار از سیستم مستندسازی PMO تولید شده است</p>
    </div>
</body>
</html>
"""
    
    return html_template


@app.route('/export/<doc_type>/<version>/<format>')
def export_document(doc_type, version, format):
    """
    Export endpoint
    Example: /export/risk-management/v3.0.0/pdf
    """
    try:
        if format not in ['pdf', 'word', 'docx']:
            return jsonify({'error': 'Invalid format. Use pdf or word/docx'}), 400
        
        # Normalize format
        if format == 'docx':
            format = 'word'
        
        cache_key = f"{doc_type}_{version}_{format}"
        
        # Check cache
        if cache_key in file_cache:
            file_path = file_cache[cache_key]
            if file_path.exists():
                print(f"✅ Serving cached file: {file_path}")
                return send_file(
                    file_path,
                    as_attachment=True,
                    download_name=file_path.name
                )
        
        # Generate new file
        print(f"🔄 Generating {format.upper()} for {doc_type} {version}...")
        
        if format == 'pdf':
            file_path = generate_pdf(doc_type, version)
        else:  # word
            file_path = generate_word(doc_type, version)
        
        if not file_path or not file_path.exists():
            return jsonify({'error': 'Failed to generate export file'}), 500
        
        # Cache the file
        file_cache[cache_key] = file_path
        
        print(f"✅ Export complete: {file_path}")
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=file_path.name
        )
        
    except Exception as e:
        print(f"❌ Export error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


def generate_pdf(doc_type, version):
    """Generate PDF using Chrome/Edge headless"""
    html_content = combine_html_for_pdf(doc_type, version)
    if not html_content:
        return None
    
    # Save temporary HTML
    temp_html = EXPORTS_DIR / f"temp_{doc_type}_{version}.html"
    with open(temp_html, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    output_file = EXPORTS_DIR / f"{doc_type}-{version}.pdf"
    
    print(f"Generating PDF: {output_file}")
    
    # Try Chrome/Edge for PDF generation
    chrome_paths = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    
    chrome_exe = None
    for path in chrome_paths:
        if Path(path).exists():
            chrome_exe = path
            break
    
    if not chrome_exe:
        print("❌ Chrome/Edge not found. Cannot generate PDF.")
        temp_html.unlink()
        return None
    
    try:
        # Use Chrome/Edge headless to generate PDF
        subprocess.run([
            chrome_exe,
            '--headless',
            '--disable-gpu',
            '--print-to-pdf=' + str(output_file.absolute()),
            str(temp_html.absolute())
        ], check=True, capture_output=True)
        
        print(f"✅ PDF generated successfully: {output_file}")
        
        # Clean up temp file
        temp_html.unlink()
        
        return output_file
    except Exception as e:
        print(f"❌ Error generating PDF: {e}")
        if temp_html.exists():
            temp_html.unlink()
        return None


def generate_word(doc_type, version):
    """Generate Word document using python-docx"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Get markdown files from docs directory
    version_path = DOCS_DIR / doc_type / 'versions' / version
    if not version_path.exists():
        print(f"Error: Version path {version_path} does not exist")
        return None
    
    md_files = sorted(version_path.glob('*.md'))
    if not md_files:
        return None
    
    # Create Word document
    doc = Document()
    
    # Set RTL direction (requires proper Word template, but we'll do our best)
    section = doc.sections[0]
    
    # Add header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(f'مستندات {doc_type.replace("-", " ")}\n')
    header_run.font.size = Pt(24)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(26, 35, 126)
    
    version_run = header_para.add_run(f'نسخه {version}')
    version_run.font.size = Pt(16)
    version_run.font.color.rgb = RGBColor(63, 81, 181)
    
    doc.add_paragraph()  # Empty line
    
    # Process each markdown file
    md = markdown.Markdown(extensions=['extra', 'tables'])
    
    for md_file in md_files:
        print(f"Processing: {md_file.name}")
        
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
            # Split by lines and process
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    continue
                
                # Process headers
                if line.startswith('# '):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(line[2:])
                    run.font.size = Pt(20)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 35, 126)
                    
                elif line.startswith('## '):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(line[3:])
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(40, 53, 147)
                    
                elif line.startswith('### '):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(line[4:])
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(48, 63, 159)
                    
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    # Numbered list
                    p = doc.add_paragraph(line[3:], style='List Number')
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                else:
                    # Regular paragraph
                    if line:
                        p = doc.add_paragraph(line)
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page break after each file (except the last one)
        if md_file != md_files[-1]:
            doc.add_page_break()
    
    # Add footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('این مستند به صورت خودکار از سیستم مستندسازی PMO تولید شده است')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    output_file = EXPORTS_DIR / f"{doc_type}-{version}.docx"
    doc.save(str(output_file))
    
    print(f"✅ Word document generated successfully: {output_file}")
    return output_file



@app.route('/health')
def health():
    """Health check endpoint"""
    return jsonify({'status': 'ok', 'service': 'export-server'})


@app.route('/')
def index():
    """Info endpoint"""
    return jsonify({
        'service': 'MkDocs Export Server',
        'endpoints': {
            'export': '/export/<doc_type>/<version>/<format>',
            'health': '/health'
        },
        'examples': [
            '/export/risk-management/v3.0.0/pdf',
            '/export/risk-management/v3.0.0/word',
            '/export/project-management/v1.0.0/pdf'
        ]
    })


def clear_cache_periodically():
    """Clear cache every 30 minutes to avoid stale files"""
    while True:
        time.sleep(1800)  # 30 minutes
        print("🗑️ Clearing export cache...")
        file_cache.clear()


if __name__ == '__main__':
    # Start cache cleanup thread
    cache_thread = threading.Thread(target=clear_cache_periodically, daemon=True)
    cache_thread.start()
    
    print("=" * 60)
    print("🚀 Export Server Starting...")
    print("=" * 60)
    print("📍 URL: http://127.0.0.1:5000")
    print("📄 Export endpoint: /export/<doc>/<version>/<format>")
    print("💡 Example: /export/risk-management/v3.0.0/pdf")
    print("=" * 60)
    
    app.run(host='127.0.0.1', port=5000, debug=True, use_reloader=False)
